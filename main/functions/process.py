import tempfile, pathlib, docx, textract
from openpyxl import Workbook
from openpyxl.styles.fills import PatternFill
from openpyxl.styles.colors import Color
from openpyxl.styles.borders import Border, Side
from openpyxl import Workbook
from docx.shared import RGBColor

from nltk import wordpunct_tokenize, pos_tag
from nltk.corpus import wordnet, stopwords

def ExtractText(filename):

    # Create the temp file
    temp = tempfile.NamedTemporaryFile(suffix='.txt', delete=False)
    pt = pathlib.Path(temp.name)

    # Get text
    text = textract.process(filename)

    with open(temp.name, 'w') as f:
        f.write(text.decode('utf-8'))
        
    return temp.name

def WritePageDocx(filename, lines):

    # create an instance of a word document
    doc = docx.Document()

    # add a heading of level 0 (largest heading)
    doc.add_heading('Paper View of the Nouns and Synonyms ', 0)
    
    paragraph = doc.add_paragraph('')

    for line in lines:

        for item in line:

            if isinstance(item, str):
                paragraph.add_run(item + ' ')

            elif isinstance(item, dict):

                for key in item:
                    if item[key] is not None:
                        lst = ' ( ' + ', '.join(item[key]) + ' ) '
                    else:
                        lst = ' '

                    run = paragraph.add_run(key + lst)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)

        paragraph.add_run('\n')

    # now save the document to a location
    doc.save(filename.name)

def WriteTableExcel(filename, nouns):

    # Make a workbook
    wb = Workbook()

    # Grab the active worksheet
    ws = wb.active

    col = 1
    next_row = 1
    counter = 0
    color = 0
    widest = [0,0]
    thin_border = Border(left=Side(style='thin'), 
                     right=Side(style='thin'), 
                     top=Side(style='thin'), 
                     bottom=Side(style='thin'))

    header = PatternFill(patternType='solid', fgColor=Color(rgb='00bb94'))
    colors = [PatternFill(patternType='solid', fgColor=Color(rgb='9BB9BA')), PatternFill(patternType='solid', fgColor=Color(rgb='B0E5EA'))]


    # Insert header
    ws.append(["#", "Noun", "Alternatives"])

    # Write fill
    ws.cell(column=1 , row=1).fill = header
    ws.cell(column=2 , row=1).fill = header
    ws.cell(column=3 , row=1).fill = header   

    # Write borders
    ws.cell(column=1 , row=1).border = thin_border
    ws.cell(column=2 , row=1).border = thin_border
    ws.cell(column=3 , row=1).border = thin_border   
    

    for key, value in nouns.items():
        next_row += 1
        counter += 1

        # Write the #
        ws.cell(column=1 , row=next_row, value=counter)

        # Write the noun
        ws.cell(column=2 , row=next_row, value=key)

        # Write the syns        
        if value is not None:
            ws.cell(column=3 , row=next_row, value=','.join(value))

        # Color in the rows
        for x in range(1, 4):
            ws.cell(column=x , row=next_row).fill = colors[color]

        # Swap color
        color = 1 - color

        # Write borders
        ws.cell(column=1 , row=next_row).border = thin_border
        ws.cell(column=2 , row=next_row).border = thin_border
        ws.cell(column=3 , row=next_row).border = thin_border

        # Find widest
        widest[0] = widest[0] if len(key) < widest[0] else len(key)
        if ws.cell(column=3 , row=next_row).value is not None:
            widest[1] = widest[1] if len(ws.cell(column=3 , row=next_row).value) < widest[1] else len(ws.cell(column=3 , row=next_row).value)

    # Change widths
    ws.column_dimensions['B'].width = widest[0]
    ws.column_dimensions['C'].width = widest[1]
        
    # Save the file
    wb.save(filename.name)

def ProcessFile(filename):
    
    # Get the data
    with open(filename, 'r', encoding='utf-8', errors='replace') as file:
        data = file.read().replace('“', "\"").replace('”', "\"").replace('’', "'").replace('–', "'")

    # Get all the words from the file and remove duplicates
    words = list(set(wordpunct_tokenize(data)))
    words = sorted([ word for word in words if word not in stopwords.words('english') and len(word) > 1 and (word.title() not in words or word.lower() not in words) ])
    
    lst = list()
    unwanted = list()

    # Check for nours and proper nouns ONLY
    for (word, tag) in pos_tag(words):

        # Check if its plural or verb
        if tag == 'NNS' or tag[0] == 'V':

            # Remove from list if it got in there
            if word in lst:
                lst.remove(word)

            # Add it to the plural list
            if word not in unwanted:
                unwanted.append(word)

        # Check if its a proper noun and not in either list
        if (tag == 'NNP' or tag == 'NN') and word not in unwanted:
            lst.append(word)

    # Convert to the nouns
    nouns = dict.fromkeys(sorted(lst))

    # Now get all the synonyms
    for noun in list(nouns):
        lst = list()

        # Get the synonyms for the noun
        for syn in wordnet.synsets(noun, pos=wordnet.NOUN):

            # Loop all the lemmas
            for lemmas in syn.lemmas():

                # Get the name
                synonym = lemmas.name()

                if '_' not in synonym and synonym not in lst and noun.lower() != synonym.lower() and (synonym.title() in nouns or synonym in nouns):
                    lst.append(synonym)

            if len(lst):
                nouns[noun] = lst

    # Copy the dictionary
    temp_nouns = dict(nouns)

    # A new list
    temp_lines = list()

    # Counter for line on
    line_counter = 0

    # Loop each line and loop for the noun
    for line in data.split('\n'):

        # Temporary sentance to add to the temp_lines
        temp_sentence = list()
        
        # Get each word to compare
        for word in line.split():

            # Check if in the temp_nouns dict
            if word in temp_nouns:

                # If so replace it
                temp_sentence.append(dict({word: nouns[word]}))

                # Remove the noun
                del temp_nouns[word]

            else:
                temp_sentence.append(word)

        # Add temp_sentence to temp_lines
        temp_lines.append(temp_sentence)

        # Increment
        line_counter += 1
        
    return temp_lines, nouns